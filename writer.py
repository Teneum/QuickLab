from docx import Document
from docx.shared import Inches
from docx.shared import Pt

from functions.configfuncs import *
from functions.read_py_file import open_file


class Writer:
    def __init__(self, ex_no, section, stud_name, date):
        self.ex_no = ex_no
        self.section = section
        self.stud_name = stud_name
        self.date = date

        self.document = Document()

        font = self.document.styles['Normal'].font
        font.name = 'Cambria'
        font.size = Pt(14)

        self.document.add_heading(f'Lab Exercise: {ex_no}', 0)
        self.document.add_paragraph(f"Date: {self.date}")
        self.document.add_paragraph(f"Name: {self.stud_name}")
        self.document.add_paragraph(f"Section: {self.section}")

    def aim(self, aim):
        self.aim = aim
        self.document.add_heading("AIM: ", 1)
        self.document.add_paragraph(aim)

    def algorithm(self, algo):
        self.document.add_heading("ALGORITHM: ", 1)
        self.document.add_paragraph(algo)

    def source_code(self):
        path = get_selected_file()
        code = open_file(path=path)
        self.document.add_heading("SOURCE CODE: ", 1)
        self.document.add_paragraph(code)

    def screenshots(self):
        self.document.add_heading("SCREENSHOTS: ", 1)
        paths = get_image_paths()
        for path in paths:
            self.document.add_picture(path, width=Inches(4), height=Inches(4))

    def sample_output(self):
        self.document.add_heading("SAMPLE OUTPUT: ", 1)

    def finish_doc(self):
        self.document.add_heading("RESULT: ", 1)
        self.document.add_paragraph("The python program has been successfully executed without error and with desired output.")
        self.document.save(f"LabEx{self.ex_no}-{(self.stud_name + self.section).strip()}.docx")
        clear_image_paths()
        clear_file_path()
        clear_images()
        delete_screenshot_dir()


